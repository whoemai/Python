import os
import tkinter as tk
from datetime import datetime
from tkinter import messagebox

import pandas as pd
from docx import Document

# Pedir ao usuário para inserir o número de dados
n = int(input("Insira o número de dados: "))
colUm = input("Digite o nome da primeira coluna: ")
colDois = input("Digite o nome da segunda coluna: ")

# Pedir ao usuário para inserir as cores e nomes
colUmData = [
    input(f"Insira o dado referente {colUm} {i+1}: ") for i in range(n)]
colDoisData = [
    input(f"Insira o dado referente {colDois} {i+1}: ") for i in range(n)]

# Criar um dataframe com as cores e nomes
dados = {colUm: colUmData, colDois: colDoisData}
df = pd.DataFrame(dados)

# Define o caminho e nome do arquivo com a data do dia
today = datetime.today().strftime('%Y-%m-%d')
filepath_excel = f'C:\\Users\\x\\Documents\\Excel\\appy-PY{today}.xlsx'
filepath_word = f'C:\\Users\\x\\Documents\\Word\\appy-PY{today}.docx'
filepath_csv = f'C:\\Users\\x\\Documents\\csv\\appy-PY{today}.csv'

# Salva os arquivos permitindo substituir os existentes
if os.path.exists(filepath_excel):
    confirm = messagebox.askyesno(
        "Arquivo existente", f"O arquivo {filepath_excel} já existe. Deseja substituí-lo?")
    if not confirm:
        messagebox.showinfo("Arquivo não salvo",
                            f"O arquivo {filepath_excel} não foi salvo.")
    else:
        df.to_excel(filepath_excel, index=False)
else:
    df.to_excel(filepath_excel, index=False)

if os.path.exists(filepath_word):
    confirm = messagebox.askyesno(
        "Arquivo existente", f"O arquivo {filepath_word} já existe. Deseja substituí-lo?")
    if not confirm:
        messagebox.showinfo("Arquivo não salvo",
                            f"O arquivo {filepath_word} não foi salvo.")
    else:
        document = Document()
        table = document.add_table(rows=1, cols=2)
        for i in range(len(df)):
            cells = table.add_row().cells
            cells[0].text = df.iloc[i][colUm]
            cells[1].text = df.iloc[i][colDois]
        document.save(filepath_word)
else:
    document = Document()
    table = document.add_table(rows=1, cols=2)
    for i in range(len(df)):
        cells = table.add_row().cells
        cells[0].text = df.iloc[i][colUm]
        cells[1].text = df.iloc[i]

print("Salvo com sucesso!")
